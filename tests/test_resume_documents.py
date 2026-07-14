import io

import pytest
from docx import Document
from PIL import Image
from pypdf import PdfReader
from reportlab.pdfgen import canvas

from resume_documents import (
    build_resume_html,
    build_resume_docx,
    build_resume_pdf,
    extract_resume_text,
    list_templates,
    normalize_resume_structure,
    parse_resume_structure,
    structured_to_markdown,
    template_preview_html,
)

RESUME_TEXT = """# 张三
AI 应用工程师 | 13800138000 | zhangsan@example.com | 深圳

## 个人简介
- 5 年 Python 后端与 AI 应用开发经验。

## 专业技能
- Python、FastAPI、PostgreSQL、RAG

## 项目经历
- 智能招聘助手：负责 Agent 工作流与检索模块，响应时间降低 30%。

## 教育经历
- 示例大学 | 计算机科学 | 本科 | 2016-2020
"""


def test_template_library_has_many_distinct_templates():
    templates = list_templates()
    assert len(templates) >= 18
    assert len({item["id"] for item in templates}) == len(templates)
    assert any(item["ats_friendly"] for item in templates)
    assert {item["layout"] for item in templates} >= {"single", "two-column", "sidebar"}
    assert all(item["features"] for item in templates)


def test_parse_resume_and_markdown_round_trip():
    structured = parse_resume_structure(RESUME_TEXT)
    assert structured["basics"]["name"] == "张三"
    assert structured["basics"]["email"] == "zhangsan@example.com"
    assert any(section["key"] == "projects" for section in structured["sections"])
    markdown = structured_to_markdown(structured)
    assert "项目经历" in markdown
    assert "FastAPI" in markdown


def test_extract_docx_upload():
    doc = Document()
    doc.add_heading("李四", level=1)
    doc.add_paragraph("Python 工程师")
    doc.add_heading("专业技能", level=2)
    doc.add_paragraph("Python / Django")
    data = io.BytesIO()
    doc.save(data)
    text, source_format = extract_resume_text("resume.docx", data.getvalue())
    assert source_format == "docx"
    assert "李四" in text
    assert "Django" in text


def test_extract_pdf_upload():
    data = io.BytesIO()
    pdf = canvas.Canvas(data)
    pdf.drawString(72, 760, "Alice Chen - Python Engineer")
    pdf.drawString(72, 740, "Skills: Python FastAPI PostgreSQL")
    pdf.save()
    text, source_format = extract_resume_text("resume.pdf", data.getvalue())
    assert source_format == "pdf"
    assert "Python Engineer" in text


def test_generate_docx_and_pdf_with_selected_template():
    structured = parse_resume_structure(RESUME_TEXT)
    docx_data = build_resume_docx(structured, "modern_blue")
    pdf_data = build_resume_pdf(structured, "modern_blue")
    assert docx_data.startswith(b"PK")
    assert pdf_data.startswith(b"%PDF")
    rendered_docx = Document(io.BytesIO(docx_data))
    assert any("张三" in paragraph.text for paragraph in rendered_docx.paragraphs)
    pdf_page = PdfReader(io.BytesIO(pdf_data)).pages[0]
    assert abs(float(pdf_page.mediabox.width) - 595.28) < 1
    assert abs(float(pdf_page.mediabox.height) - 841.89) < 1
    docx_section = rendered_docx.sections[0]
    assert abs(docx_section.page_width.mm - 210) < 0.2
    assert abs(docx_section.page_height.mm - 297) < 0.2


def test_template_preview_is_safe_html():
    result = template_preview_html("ats_classic", parse_resume_structure("# <script>x</script>\n## 技能\n- Python"))
    assert "<script>" not in result
    assert "&lt;script&gt;" in result


def test_layout_templates_render_real_structures_with_resume_content():
    structured = parse_resume_structure(RESUME_TEXT)
    sidebar = template_preview_html("vivi_sidebar", structured)
    columns = template_preview_html("vivi_two_column", structured)
    timeline = template_preview_html("vivi_classic", structured)
    assert "张三" in sidebar and "layout-sidebar" in sidebar and "resume-aside" in sidebar
    assert "layout-two-column" in columns and columns.count("resume-column") >= 2
    assert "has-timeline" in timeline


def test_html_export_contains_print_ready_template_css():
    result = build_resume_html(parse_resume_structure(RESUME_TEXT), "vivi_sidebar")
    assert "@page{size:A4" in result
    assert "layout-sidebar" in result
    assert "resume-shell" in result


def test_new_visual_templates_export_docx_and_pdf():
    structured = parse_resume_structure(RESUME_TEXT)
    for template_id in ("vivi_classic", "vivi_modern", "vivi_two_column", "vivi_sidebar", "vivi_professional"):
        assert build_resume_docx(structured, template_id).startswith(b"PK")
        assert build_resume_pdf(structured, template_id).startswith(b"%PDF")


def test_two_column_docx_uses_distinct_side_and_main_columns():
    doc = Document(io.BytesIO(build_resume_docx(parse_resume_structure(RESUME_TEXT), "vivi_two_column")))
    assert any(len(table.columns) == 2 for table in doc.tables)


def test_work_descriptions_preserve_bullet_and_numbered_lists():
    structured = normalize_resume_structure({
        "basics": {"name": "周宁", "title": "产品经理"},
        "sections": [{
            "key": "experience",
            "title": "工作经历",
            "entries": [{
                "company": "示例科技",
                "role": "产品经理",
                "description": "• 推动核心流程上线\n1. 建立指标体系\n2. 转化率提升 18%",
            }],
        }],
    })
    preview = template_preview_html("vivi_elegant", structured)
    assert "<span>•</span><span>推动核心流程上线</span>" in preview
    assert "<span>1.</span><span>建立指标体系</span>" in preview
    doc = Document(io.BytesIO(build_resume_docx(structured, "vivi_elegant")))
    styles = [paragraph.style.name for paragraph in doc.paragraphs]
    assert "List Bullet" in styles
    assert "List Number" in styles
    assert build_resume_pdf(structured, "vivi_elegant").startswith(b"%PDF")


def test_structured_editor_schema_preserves_order_visibility_and_style():
    structured = normalize_resume_structure({
        "basics": {"name": "林微", "title": "产品设计师"},
        "section_order": ["basic", "education", "summary", "experience", "projects", "skills", "evaluation"],
        "hidden_sections": ["evaluation"],
        "style": {"accent_color": "#0F766E", "font_family": "DengXian", "font_size": 14, "line_height": 1.7, "page_padding": 48},
        "sections": [
            {"key": "education", "title": "教育经历", "entries": [{"school": "示例大学", "major": "工业设计", "degree": "本科"}]},
            {"key": "summary", "title": "个人简介", "content": "关注 AI 产品体验。"},
            {"key": "evaluation", "title": "自我评价", "content": "认真负责。"},
        ],
    })
    preview = template_preview_html("vivi_elegant", structured)
    assert structured["schema_version"] == 2
    assert structured["section_order"][1:3] == ["education", "summary"]
    assert "--accent:#0F766E" in preview
    assert "DengXian" in preview
    assert "示例大学" in preview
    assert "自我评价" not in preview
    assert preview.index("教育经历") < preview.index("个人简介")


def test_field_level_font_size_and_line_height_are_preserved_in_exports():
    structured = normalize_resume_structure({
        "basics": {
            "name": "字段字号测试",
            "title": "产品经理",
            "phone": "13800138000",
            "field_styles": {
                "name": {"font_size": 30, "line_height": 1.2},
                "title": {"font_size": 15, "line_height": 1.4},
                "phone": {"font_size": 40, "line_height": 2.4},
            },
        },
        "style": {"page_padding": 48, "section_spacing": 20},
        "sections": [
            {"key": "summary", "title": "个人简介", "content": "字段级正文样式。", "field_styles": {"content": {"font_size": 12, "line_height": 1.8}}},
            {"key": "experience", "title": "工作经历", "entries": [{
                "company": "示例公司",
                "role": "负责人",
                "description": "负责核心产品。",
                "field_styles": {
                    "company": {"font_size": 14, "line_height": 1.3},
                    "role": {"font_size": 11, "line_height": 1.25},
                    "description": {"font_size": 10, "line_height": 1.7},
                    "phone": {"font_size": 40},
                },
            }]},
        ],
    })
    assert structured["basics"]["field_styles"]["name"] == {"font_size": 30.0, "line_height": 1.2}
    assert "phone" not in structured["basics"]["field_styles"]
    experience = next(section for section in structured["sections"] if section["key"] == "experience")
    assert "phone" not in experience["entries"][0]["field_styles"]

    preview = template_preview_html("vivi_elegant", structured)
    assert '<h1 style="font-size:30px;line-height:1.2">字段字号测试</h1>' in preview
    assert '<div class="section-copy" style="font-size:12px;line-height:1.8">' in preview
    assert '--page-padding:48px;--section-spacing:20px' in preview

    docx = Document(io.BytesIO(build_resume_docx(structured, "vivi_elegant")))
    name_paragraph = next(paragraph for paragraph in docx.paragraphs if paragraph.text == "字段字号测试")
    summary_paragraph = next(paragraph for paragraph in docx.paragraphs if paragraph.text == "字段级正文样式。")
    assert abs(name_paragraph.runs[0].font.size.pt - 22.8) < 0.35
    assert abs(float(name_paragraph.paragraph_format.line_spacing) - 1.2) < 0.01
    assert abs(summary_paragraph.runs[0].font.size.pt - 9.12) < 0.35
    assert abs(float(summary_paragraph.paragraph_format.line_spacing) - 1.8) < 0.01
    assert build_resume_pdf(structured, "vivi_elegant").startswith(b"%PDF")


def test_hidden_personal_profile_is_omitted_from_html_and_docx():
    structured = normalize_resume_structure({
        "basics": {"name": "不应显示的姓名", "title": "测试岗位"},
        "hidden_sections": ["basic"],
        "sections": [{"key": "summary", "title": "个人简介", "content": "只显示正文。"}],
    })
    html_result = template_preview_html("modern_blue", structured)
    docx_result = Document(io.BytesIO(build_resume_docx(structured, "modern_blue")))
    assert "不应显示的姓名" not in html_result
    assert not any("不应显示的姓名" in paragraph.text for paragraph in docx_result.paragraphs)


def test_photo_age_and_centered_profile_render_in_every_template():
    photo = io.BytesIO()
    Image.new("RGB", (300, 400), "#7DD3FC").save(photo, format="PNG")
    import base64

    structured = normalize_resume_structure({
        "basics": {
            "name": "林微",
            "title": "产品经理",
            "age": "29",
            "photo": "data:image/png;base64," + base64.b64encode(photo.getvalue()).decode("ascii"),
        },
    })
    for template in list_templates():
        preview = template_preview_html(template["id"], structured)
        assert "resume-photo" in preview
        assert 'alt="个人照片"' in preview
        assert "29 岁" in preview
        assert "align-center" in preview or template["layout"] == "sidebar"
    docx = Document(io.BytesIO(build_resume_docx(structured, "vivi_elegant")))
    assert len(docx.inline_shapes) >= 1
    assert build_resume_pdf(structured, "vivi_sidebar").startswith(b"%PDF")


def test_birthday_migrates_to_age_and_invalid_date_range_is_rejected():
    migrated = normalize_resume_structure({"basics": {"name": "旧简历", "birthday": "1998-06-01"}})
    assert migrated["basics"]["age"].isdigit()
    assert "birthday" not in migrated["basics"]
    with pytest.raises(ValueError, match="结束时间不能早于开始时间"):
        normalize_resume_structure({
            "basics": {"name": "时间错误"},
            "sections": [{
                "key": "experience",
                "title": "工作经历",
                "entries": [{"company": "示例公司", "start_date": "2025.06", "end_date": "2024.12"}],
            }],
        })
