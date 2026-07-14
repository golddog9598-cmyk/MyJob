"""简历上传解析、结构化和模板化 DOCX/PDF/HTML 生成。"""

from __future__ import annotations

import html
import io
import json
import re
import zipfile
import base64
import binascii
from copy import deepcopy
from datetime import date
from pathlib import Path
from typing import Any

MAX_UPLOAD_BYTES = 15 * 1024 * 1024
SUPPORTED_EXTENSIONS = {".docx", ".pdf", ".txt", ".md", ".markdown", ".html", ".htm", ".rtf", ".json", ".odt"}


# 模板配置参考 vivi-resume 的“内容与视觉配置分离”思路：同一份结构化简历可切换
# 头部、栏目、时间线、双栏/侧栏和密度，不改写用户内容。旧字段继续保留，避免破坏
# 已保存的 template_id 及既有 DOCX/PDF 导出链接。
RESUME_TEMPLATES: list[dict[str, Any]] = [
    {"id": "ats_classic", "name": "ATS 经典", "category": "ATS 友好", "description": "黑白单栏，机器解析稳定，适合绝大多数岗位", "accent": "#202124", "muted": "#5F6368", "font": "Microsoft YaHei", "header": "center", "header_style": "plain", "section": "rule", "layout": "single", "ats_friendly": True},
    {"id": "modern_blue", "name": "现代蓝", "category": "通用", "description": "清爽蓝色标题栏，适合产品、运营与互联网岗位", "accent": "#2563EB", "muted": "#64748B", "font": "Microsoft YaHei", "header": "left", "header_style": "band", "section": "rule", "layout": "single", "ats_friendly": True},
    {"id": "tech_slate", "name": "技术灰蓝", "category": "技术", "description": "克制的工程感，突出技术栈、项目和交付成果", "accent": "#334155", "muted": "#64748B", "font": "Microsoft YaHei", "header": "left", "header_style": "plain", "section": "band", "layout": "single", "ats_friendly": True},
    {"id": "minimal_gray", "name": "极简灰", "category": "极简", "description": "低干扰、高留白，适合咨询、法务和职能岗位", "accent": "#374151", "muted": "#9CA3AF", "font": "Microsoft YaHei", "header": "center", "header_style": "plain", "section": "plain", "layout": "single", "ats_friendly": True},
    {"id": "executive_navy", "name": "高管藏青", "category": "管理", "description": "稳重藏青横幅，适合管理与资深岗位", "accent": "#0F2A44", "muted": "#52677D", "font": "Microsoft YaHei", "header": "left", "header_style": "band", "section": "band", "layout": "single", "ats_friendly": True},
    {"id": "product_orange", "name": "产品橙", "category": "产品", "description": "温和强调色，突出成果、指标与业务影响", "accent": "#C2410C", "muted": "#78716C", "font": "Microsoft YaHei", "header": "left", "header_style": "plain", "section": "rule", "layout": "single", "ats_friendly": True},
    {"id": "growth_green", "name": "增长绿", "category": "运营", "description": "适合增长、市场和商业化岗位", "accent": "#047857", "muted": "#6B7280", "font": "Microsoft YaHei", "header": "left", "header_style": "band", "section": "band", "layout": "single", "ats_friendly": True},
    {"id": "data_purple", "name": "数据紫", "category": "数据", "description": "突出分析、研究和量化项目", "accent": "#6D28D9", "muted": "#6B7280", "font": "Microsoft YaHei", "header": "center", "header_style": "plain", "section": "rule", "layout": "single", "timeline": True, "ats_friendly": True},
    {"id": "academic_serif", "name": "学术雅宋", "category": "学术", "description": "宋体排版，适合研究、教育与应届生", "accent": "#7C2D12", "muted": "#78716C", "font": "SimSun", "header": "center", "header_style": "plain", "section": "plain", "layout": "single", "ats_friendly": True},
    {"id": "compact_one_page", "name": "一页紧凑", "category": "一页式", "description": "压缩字号和间距，适合 3-5 年经验候选人", "accent": "#1D4ED8", "muted": "#64748B", "font": "Microsoft YaHei", "header": "inline", "header_style": "inline", "section": "rule", "layout": "single", "compact": True, "ats_friendly": True},
    {"id": "graduate_fresh", "name": "校招清新", "category": "应届生", "description": "教育和项目优先，适合校招与实习", "accent": "#0E7490", "muted": "#64748B", "font": "Microsoft YaHei", "header": "center", "header_style": "band", "section": "band", "layout": "single", "education_first": True, "ats_friendly": True},
    {"id": "creative_rose", "name": "创意玫红", "category": "设计创意", "description": "有辨识度的色块设计，适合设计与内容岗位", "accent": "#BE185D", "muted": "#6B7280", "font": "Microsoft YaHei", "header": "left", "header_style": "band", "section": "band", "layout": "single", "ats_friendly": False},
    {"id": "vivi_classic", "name": "经典时间线", "category": "经典", "description": "居中头部与时间线组合，职业历程清晰", "accent": "#7C5CFC", "muted": "#4A4A6A", "font": "Microsoft YaHei", "header": "center", "header_style": "plain", "section": "rule", "layout": "single", "timeline": True, "ats_friendly": True},
    {"id": "vivi_modern", "name": "现代青", "category": "现代", "description": "彩色头部与卡片栏目，年轻、有活力", "accent": "#0891B2", "muted": "#64748B", "font": "Microsoft YaHei", "header": "left", "header_style": "band", "section": "card", "layout": "single", "ats_friendly": False},
    {"id": "vivi_two_column", "name": "信息双栏", "category": "双栏", "description": "左侧能力与教育，右侧经历与项目，信息密度高", "accent": "#2563EB", "muted": "#64748B", "font": "Microsoft YaHei", "header": "left", "header_style": "band", "section": "rule", "layout": "two-column", "ats_friendly": False},
    {"id": "vivi_sidebar", "name": "蓝色侧栏", "category": "侧栏", "description": "独立个人信息侧栏，视觉层次鲜明", "accent": "#3B82F6", "muted": "#64748B", "font": "Microsoft YaHei", "header": "left", "header_style": "sidebar", "section": "plain", "layout": "sidebar", "sidebar_bg": "#DBEAFE", "ats_friendly": False},
    {"id": "vivi_professional", "name": "专业黑金", "category": "管理", "description": "深色标题条与简洁正文，稳重干练", "accent": "#1F2937", "muted": "#6B7280", "font": "Microsoft YaHei", "header": "left", "header_style": "band", "section": "dark", "layout": "single", "ats_friendly": True},
    {"id": "vivi_elegant", "name": "优雅翡翠", "category": "通用", "description": "细腻留白与翡翠绿细节，专业而不沉闷", "accent": "#059669", "muted": "#6B7280", "font": "Microsoft YaHei", "header": "left", "header_style": "plain", "section": "rule", "layout": "single", "ats_friendly": True},
]


SECTION_ALIASES = {
    "summary": {"个人简介", "职业简介", "职业概述", "求职概述", "summary", "profile", "about"},
    "skills": {"专业技能", "技能", "技能清单", "核心能力", "技术栈", "skills", "competencies"},
    "experience": {"工作经历", "工作经验", "职业经历", "实习经历", "experience", "employment", "work experience"},
    "projects": {"项目经历", "项目经验", "代表项目", "projects", "project experience"},
    "education": {"教育经历", "教育背景", "学历", "education"},
    "evaluation": {"自我评价", "个人评价", "自我描述", "evaluation", "self evaluation"},
    "certifications": {"证书", "资质证书", "认证", "荣誉奖项", "certifications", "awards"},
    "other": {"其他", "补充信息", "语言能力", "作品集", "other"},
}

EDITOR_SECTION_TITLES = {
    "basic": "个人资料",
    "summary": "个人简介",
    "education": "教育经历",
    "experience": "工作经历",
    "projects": "项目经历",
    "skills": "专业技能",
    "evaluation": "自我评价",
}
EDITOR_SECTION_ORDER = ["basic", "summary", "experience", "education", "projects", "skills", "evaluation"]
SECTION_TEXT_FIELDS = {"summary": {"content"}, "evaluation": {"content"}}
ENTRY_TEXT_FIELDS = {
    "experience": {"company", "role", "description"},
    "education": {"school", "major", "degree", "description"},
    "projects": {"name", "role", "technologies", "description"},
    "skills": {"content"},
}
MAX_PHOTO_BYTES = 3 * 1024 * 1024


def _normalize_photo(value: Any) -> str:
    photo = str(value or "").strip()
    if not photo:
        return ""
    match = re.fullmatch(r"data:image/(jpeg|jpg|png|webp);base64,([A-Za-z0-9+/=\r\n]+)", photo, re.IGNORECASE)
    if not match:
        raise ValueError("照片格式无效，仅支持 JPG、PNG 或 WebP")
    try:
        payload = base64.b64decode(match.group(2), validate=True)
    except (binascii.Error, ValueError) as exc:
        raise ValueError("照片数据损坏，请重新上传") from exc
    if len(payload) > MAX_PHOTO_BYTES:
        raise ValueError("压缩后的照片不能超过 3MB")
    try:
        from PIL import Image

        with Image.open(io.BytesIO(payload)) as image:
            image.verify()
    except Exception as exc:
        raise ValueError("照片内容无效，请重新上传") from exc
    subtype = "jpeg" if match.group(1).lower() in {"jpg", "jpeg"} else match.group(1).lower()
    return f"data:image/{subtype};base64,{base64.b64encode(payload).decode('ascii')}"


def _normalize_field_styles(value: Any, allowed_fields: set[str]) -> dict[str, dict[str, float]]:
    if not isinstance(value, dict):
        return {}
    result = {}
    for field in allowed_fields:
        source = value.get(field)
        if not isinstance(source, dict):
            continue
        style = {}
        if source.get("font_size") not in (None, ""):
            try:
                style["font_size"] = max(8.0, min(40.0, float(source["font_size"])))
            except (TypeError, ValueError):
                pass
        if source.get("line_height") not in (None, ""):
            try:
                style["line_height"] = max(1.0, min(2.4, float(source["line_height"])))
            except (TypeError, ValueError):
                pass
        if style:
            result[field] = style
    return result


def _field_style_values(owner: dict, key: str, default_font_size: float, default_line_height: float) -> tuple[float, float]:
    style = (owner.get("field_styles") or {}).get(key) or {}
    return float(style.get("font_size") or default_font_size), float(style.get("line_height") or default_line_height)


def _field_css(owner: dict, key: str) -> str:
    style = (owner.get("field_styles") or {}).get(key) or {}
    declarations = []
    if style.get("font_size"):
        declarations.append(f"font-size:{float(style['font_size']):g}px")
    if style.get("line_height"):
        declarations.append(f"line-height:{float(style['line_height']):g}")
    return f' style="{";".join(declarations)}"' if declarations else ""


def _month_value(value: Any) -> float | None:
    normalized = str(value or "").strip()
    if re.fullmatch(r"至今|现在|present", normalized, re.IGNORECASE):
        return float("inf")
    match = re.match(r"^(\d{4})[.\-/年](\d{1,2})", normalized)
    if not match:
        return None
    month = int(match.group(2))
    return int(match.group(1)) * 12 + month if 1 <= month <= 12 else None


def _validate_entry_dates(sections: list[dict]) -> None:
    for section in sections:
        if section.get("key") not in {"experience", "education", "projects"}:
            continue
        for index, entry in enumerate(section.get("entries") or [], start=1):
            start = _month_value(entry.get("start_date") or entry.get("startDate"))
            end = _month_value(entry.get("end_date") or entry.get("endDate"))
            if start is not None and end is not None and end < start:
                title = section.get("title") or EDITOR_SECTION_TITLES.get(section.get("key"), "经历")
                raise ValueError(f"{title}第 {index} 条的结束时间不能早于开始时间")


def get_template(template_id: str | None) -> dict:
    wanted = template_id or "ats_classic"
    for item in RESUME_TEMPLATES:
        if item["id"] == wanted:
            return deepcopy(item)
    raise ValueError(f"未知简历模板: {wanted}")


def list_templates() -> list[dict]:
    templates = []
    layout_labels = {"single": "单栏", "two-column": "双栏", "sidebar": "侧栏"}
    for source in RESUME_TEMPLATES:
        item = deepcopy(source)
        item["layout"] = item.get("layout", "single")
        item["layout_label"] = layout_labels.get(item["layout"], "单栏")
        features = [item["layout_label"]]
        if item.get("timeline"):
            features.append("时间线")
        if item.get("compact"):
            features.append("紧凑")
        if item.get("ats_friendly"):
            features.append("ATS")
        item["features"] = features
        templates.append(item)
    return templates


def normalize_resume_structure(structured: dict | None) -> dict:
    """把上传解析稿和 Vue 编辑器稿统一为可持久化的 v2 结构。"""
    def number(value, default, minimum, maximum, integer=False):
        try:
            parsed = float(value)
        except (TypeError, ValueError):
            parsed = float(default)
        parsed = max(minimum, min(maximum, parsed))
        return int(parsed) if integer else parsed

    source = deepcopy(structured or {})
    basics_source = source.get("basics") if isinstance(source.get("basics"), dict) else {}
    basics = {
        key: str(basics_source.get(key) or "").strip()
        for key in ("name", "title", "phone", "email", "location", "url", "wechat", "age")
    }
    if not basics["age"]:
        birthday_match = re.match(r"^(\d{4})", str(basics_source.get("birthday") or "").strip())
        if birthday_match:
            basics["age"] = str(max(0, date.today().year - int(birthday_match.group(1))))
    if basics["age"]:
        if not basics["age"].isdigit() or not 16 <= int(basics["age"]) <= 100:
            raise ValueError("年龄必须是 16 到 100 之间的整数")
    basics["photo"] = _normalize_photo(basics_source.get("photo"))
    basics["field_styles"] = _normalize_field_styles(basics_source.get("field_styles"), {"name", "title"})
    basics["name"] = basics["name"] or "个人简历"
    raw_sections = source.get("sections") if isinstance(source.get("sections"), list) else []
    sections = []
    seen: set[str] = set()
    for index, raw in enumerate(raw_sections):
        if not isinstance(raw, dict):
            continue
        section = deepcopy(raw)
        key = str(section.get("key") or "other").strip().lower()
        if key == "work":
            key = "experience"
        if key in seen and key in EDITOR_SECTION_TITLES:
            existing = next(item for item in sections if item["key"] == key)
            existing.setdefault("items", []).extend(section.get("items") or [])
            existing.setdefault("entries", []).extend(section.get("entries") or [])
            continue
        seen.add(key)
        section["key"] = key
        section["title"] = str(section.get("title") or EDITOR_SECTION_TITLES.get(key) or "其他").strip()
        section["id"] = str(section.get("id") or f"section-{key}-{index}")
        section["items"] = [str(item).strip() for item in (section.get("items") or []) if str(item).strip()]
        section["entries"] = [entry for entry in (section.get("entries") or []) if isinstance(entry, dict)]
        if section.get("content") is not None:
            section["content"] = str(section.get("content") or "").strip()
        sections.append(section)
    for key in EDITOR_SECTION_ORDER[1:]:
        if key not in seen:
            sections.append({"id": f"section-{key}", "key": key, "title": EDITOR_SECTION_TITLES[key], "items": [], "entries": [], "content": ""})
    for section in sections:
        key = section.get("key")
        section["field_styles"] = _normalize_field_styles(section.get("field_styles"), SECTION_TEXT_FIELDS.get(key, set()))
        normalized_entries = []
        for entry in section.get("entries") or []:
            normalized_entry = deepcopy(entry)
            normalized_entry["field_styles"] = _normalize_field_styles(normalized_entry.get("field_styles"), ENTRY_TEXT_FIELDS.get(key, set()))
            normalized_entries.append(normalized_entry)
        section["entries"] = normalized_entries
    _validate_entry_dates(sections)
    requested_order = ["experience" if key == "work" else str(key) for key in (source.get("section_order") or [])]
    section_order = []
    for key in requested_order + EDITOR_SECTION_ORDER:
        if key in EDITOR_SECTION_ORDER and key not in section_order:
            section_order.append(key)
    hidden_sections = []
    for key in source.get("hidden_sections") or []:
        normalized_key = "experience" if key == "work" else str(key)
        if normalized_key in EDITOR_SECTION_ORDER and normalized_key not in hidden_sections:
            hidden_sections.append(normalized_key)
    style_source = source.get("style") if isinstance(source.get("style"), dict) else {}
    accent_color = str(style_source.get("accent_color") or "").strip()
    if accent_color and not re.fullmatch(r"#[0-9A-Fa-f]{6}", accent_color):
        accent_color = ""
    allowed_fonts = {"Microsoft YaHei", "SimSun", "DengXian", "Arial", "Noto Sans SC"}
    font_family = str(style_source.get("font_family") or "Microsoft YaHei").strip()
    if font_family not in allowed_fonts:
        font_family = "Microsoft YaHei"
    style = {
        "accent_color": accent_color,
        "font_family": font_family,
        "font_size": number(style_source.get("font_size"), 13, 10, 16),
        "line_height": number(style_source.get("line_height"), 1.55, 1.2, 2.0),
        "page_padding": number(style_source.get("page_padding"), 42, 24, 72, True),
        "section_spacing": number(style_source.get("section_spacing"), 15, 8, 28, True),
    }
    return {
        "schema_version": 2,
        "basics": basics,
        "sections": sections,
        "section_order": section_order,
        "hidden_sections": hidden_sections,
        "style": style,
        "source_text": str(source.get("source_text") or ""),
    }


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                lines.append(" | ".join(values))
    return "\n".join(lines)


def _extract_pdf(data: bytes) -> str:
    import pdfplumber

    pages = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = (page.extract_text(x_tolerance=2, y_tolerance=3) or "").strip()
            if text:
                pages.append(text)
    if not pages:
        raise ValueError("PDF 未提取到文本；扫描件请先进行 OCR")
    return "\n\n".join(pages)


def _extract_html(data: bytes) -> str:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(_decode_text(data), "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()
    return "\n".join(line.strip() for line in soup.get_text("\n").splitlines() if line.strip())


def _extract_rtf(data: bytes) -> str:
    raw = _decode_text(data)
    raw = re.sub(r"\\par[d]?\b", "\n", raw)
    raw = re.sub(r"\\u(-?\d+)\??", lambda m: chr(int(m.group(1)) % 65536), raw)
    raw = re.sub(r"\\'[0-9a-fA-F]{2}", "", raw)
    raw = re.sub(r"\\[a-zA-Z]+-?\d* ?", "", raw)
    raw = raw.replace("{", "").replace("}", "")
    return "\n".join(line.strip() for line in raw.splitlines() if line.strip())


def _extract_odt(data: bytes) -> str:
    from bs4 import BeautifulSoup

    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        xml = archive.read("content.xml")
    soup = BeautifulSoup(xml, "xml")
    return "\n".join(node.get_text(" ", strip=True) for node in soup.find_all(["text:p", "text:h"]))


def extract_resume_text(filename: str, data: bytes) -> tuple[str, str]:
    if not data:
        raise ValueError("上传文件为空")
    if len(data) > MAX_UPLOAD_BYTES:
        raise ValueError("简历文件不能超过 15MB")
    ext = Path(filename or "resume.txt").suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError("支持 DOCX、PDF、TXT、Markdown、HTML、RTF、JSON、ODT 格式")
    if ext == ".docx":
        text = _extract_docx(data)
    elif ext == ".pdf":
        text = _extract_pdf(data)
    elif ext in {".html", ".htm"}:
        text = _extract_html(data)
    elif ext == ".rtf":
        text = _extract_rtf(data)
    elif ext == ".odt":
        text = _extract_odt(data)
    elif ext == ".json":
        parsed = json.loads(_decode_text(data))
        if isinstance(parsed, dict) and parsed.get("sections"):
            return structured_to_markdown(parsed), ext.lstrip(".")
        text = json.dumps(parsed, ensure_ascii=False, indent=2) if not isinstance(parsed, str) else parsed
    else:
        text = _decode_text(data)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if len(text) < 20:
        raise ValueError("简历解析后的有效文本过少")
    return text, ext.lstrip(".")


def _canonical_section(title: str) -> str:
    cleaned = re.sub(r"[#：:\s]", "", title or "").lower()
    for key, aliases in SECTION_ALIASES.items():
        if cleaned in {re.sub(r"\s", "", x.lower()) for x in aliases}:
            return key
    return "other"


def parse_resume_structure(text: str) -> dict:
    """把自由文本保守地拆成结构，原文内容不会被模型改写。"""
    raw_lines = [line.strip() for line in (text or "").replace("\r", "").split("\n") if line.strip()]
    if not raw_lines:
        return {"basics": {}, "sections": []}
    email = re.search(r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}", text)
    phone = re.search(r"(?<!\d)(?:\+?86[- ]?)?1[3-9]\d{9}(?!\d)", text)
    url = re.search(r"https?://[^\s|]+", text)
    first = re.sub(r"^#+\s*", "", raw_lines[0]).strip()
    name = (
        first
        if 1 < len(first) <= 20
        and not re.search(r"@|1[3-9]\d{9}|简历|resume", first, re.I)
        else "个人简历"
    )
    basics = {
        "name": name,
        "title": "",
        "email": email.group(0) if email else "",
        "phone": phone.group(0) if phone else "",
        "location": "",
        "url": url.group(0) if url else "",
    }
    heading_lookup = {re.sub(r"\s", "", alias.lower()) for aliases in SECTION_ALIASES.values() for alias in aliases}
    sections: list[dict] = []
    current = {"key": "summary", "title": "个人简介", "items": []}

    for index, line in enumerate(raw_lines):
        clean = re.sub(r"^#{1,6}\s*", "", line).strip().strip("：:")
        normalized = re.sub(r"\s", "", clean.lower())
        is_heading = line.startswith("#") or normalized in heading_lookup
        if index == 0 and basics["name"] != "个人简历":
            continue
        if is_heading and len(clean) <= 20:
            if current["items"]:
                sections.append(current)
            current = {"key": _canonical_section(clean), "title": clean, "items": []}
            continue
        if index < 6 and not basics["title"] and line not in {basics["email"], basics["phone"]}:
            if any(word in line.lower() for word in ("工程师", "经理", "设计", "运营", "开发", "产品", "analyst", "engineer", "manager")):
                parts = [part.strip() for part in re.split(r"[|｜]", line) if part.strip()]
                basics["title"] = re.sub(r"^[-*•]\s*", "", parts[0] if parts else line)
                for part in reversed(parts[1:]):
                    if part not in {basics["email"], basics["phone"], basics["url"]} and not re.search(r"@|1[3-9]\d{9}", part):
                        basics["location"] = part
                        break
                continue
        item = re.sub(r"^[-*•]\s*", "", line).strip()
        if item and item not in {basics["email"], basics["phone"], basics["url"]}:
            current["items"].append(item)
    if current["items"]:
        sections.append(current)
    if not sections:
        sections = [{"key": "summary", "title": "个人简介", "items": raw_lines[1:]}]
    if any(s["key"] == "education" for s in sections):
        pass
    return normalize_resume_structure({"basics": basics, "sections": sections, "source_text": text.strip()})


def _date_range(entry: dict) -> str:
    start = str(entry.get("start_date") or entry.get("startDate") or "").strip()
    end = str(entry.get("end_date") or entry.get("endDate") or "").strip()
    return " - ".join(value for value in (start, end) if value)


def _description_items(value: str) -> list[tuple[str, str, str]]:
    """Return (kind, marker, text) while accepting plain, bullet and numbered lines."""
    items = []
    for line in str(value or "").splitlines():
        raw = line.strip()
        if not raw:
            continue
        ordered = re.match(r"^(\d+)[.)、]\s*(.*)$", raw)
        if ordered:
            text = ordered.group(2).strip()
            if text:
                items.append(("number", f"{ordered.group(1)}.", text))
            continue
        text = re.sub(r"^[-*•]\s*", "", raw).strip()
        if text:
            items.append(("bullet", "•", text))
    return items


def _photo_bytes(basics: dict) -> bytes | None:
    photo = str(basics.get("photo") or "")
    if not photo or "," not in photo:
        return None


def _framed_photo_bytes(basics: dict) -> bytes:
    from PIL import Image, ImageDraw, ImageOps

    raw = _photo_bytes(basics)
    if raw:
        try:
            with Image.open(io.BytesIO(raw)) as source:
                image = ImageOps.fit(source.convert("RGB"), (600, 800), method=Image.Resampling.LANCZOS)
        except Exception:
            raw = None
    if not raw:
        image = Image.new("RGB", (600, 800), "#F3F4F6")
        draw = ImageDraw.Draw(image)
        draw.ellipse((205, 185, 395, 375), fill="#CBD5E1")
        draw.rounded_rectangle((125, 410, 475, 680), radius=110, fill="#CBD5E1")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((5, 5, 594, 794), radius=34, outline="#94A3B8", width=10)
    output = io.BytesIO()
    image.save(output, format="PNG", optimize=True)
    return output.getvalue()
    try:
        return base64.b64decode(photo.split(",", 1)[1], validate=True)
    except (binascii.Error, ValueError):
        return None


def _contact_values(basics: dict) -> list[str]:
    values = [str(basics.get(key)) for key in ("phone", "email", "location", "url") if basics.get(key)]
    if basics.get("age"):
        values.append(f"{basics['age']} 岁")
    if basics.get("wechat"):
        values.append(f"微信：{basics['wechat']}")
    return values


def _photo_html(basics: dict) -> str:
    photo = str(basics.get("photo") or "")
    content = f'<img src="{html.escape(photo, quote=True)}" alt="个人照片">' if photo else "<span>照片</span>"
    return f'<div class="resume-photo">{content}</div>'


def _section_display_items(section: dict) -> list[str]:
    key = section.get("key")
    items = [str(item).strip() for item in (section.get("items") or []) if str(item).strip()]
    content = str(section.get("content") or "").strip()
    if content:
        items.extend(line.strip("-• ") for line in content.splitlines() if line.strip())
    for entry in section.get("entries") or []:
        if not isinstance(entry, dict):
            continue
        dates = _date_range(entry)
        if key == "education":
            headline = " | ".join(str(entry.get(field) or "").strip() for field in ("school", "major", "degree") if entry.get(field))
        elif key == "experience":
            headline = " | ".join(str(entry.get(field) or "").strip() for field in ("company", "role") if entry.get(field))
        elif key == "projects":
            headline = " | ".join(str(entry.get(field) or "").strip() for field in ("name", "role") if entry.get(field))
        else:
            headline = str(entry.get("content") or entry.get("name") or "").strip()
        if dates:
            headline = " | ".join(value for value in (headline, dates) if value)
        if headline:
            items.append(headline)
        technologies = entry.get("technologies") or []
        if isinstance(technologies, str):
            technologies = [technologies]
        if technologies:
            items.append("技术：" + " / ".join(str(value).strip() for value in technologies if str(value).strip()))
        items.extend(text for _, _, text in _description_items(entry.get("description") or ""))
    return items


def structured_to_markdown(structured: dict) -> str:
    data = normalize_resume_structure(structured)
    basics = data.get("basics") or {}
    lines = [f"# {basics.get('name') or '个人简历'}"]
    if basics.get("title"):
        lines.append(str(basics["title"]))
    contact = _contact_values(basics)
    if contact:
        lines.append(" | ".join(contact))
    for section in _ordered_sections(data, {}):
        lines.extend(["", f"## {section.get('title') or '其他'}"])
        lines.extend(f"- {item}" for item in _section_display_items(section))
    return "\n".join(lines).strip()


def _ordered_sections(structured: dict, template: dict) -> list[dict]:
    sections = list(structured.get("sections") or [])
    hidden = set(structured.get("hidden_sections") or [])
    sections = [section for section in sections if section.get("key") not in hidden and _section_display_items(section)]
    requested = ["experience" if key == "work" else key for key in (structured.get("section_order") or [])]
    if requested:
        order = {key: index for index, key in enumerate(requested)}
        original_index = {id(item): index for index, item in enumerate(sections)}
        sections.sort(key=lambda item: order.get(item.get("key"), len(order) + original_index[id(item)]))
    if template.get("education_first"):
        priority = {"summary": 0, "education": 1, "skills": 2, "projects": 3, "experience": 4}
        sections.sort(key=lambda item: priority.get(item.get("key"), 9))
    return sections


def resume_preview_css() -> str:
    """HTML 导出和控制台预览共用的模板样式。"""
    return """
.resume-preview{--accent:#2563eb;--muted:#64748b;--sidebar:#dbeafe;--resume-font:'Microsoft YaHei';--resume-font-size:13px;--resume-line-height:1.55;--page-padding:42px;--section-spacing:15px;box-sizing:border-box;background:#fff;color:#1f2937;width:100%;min-height:100%;font-family:var(--resume-font),Arial,sans-serif;font-size:var(--resume-font-size);line-height:var(--resume-line-height);overflow:hidden;text-rendering:optimizeLegibility}
.resume-preview *{box-sizing:border-box}.resume-preview header{padding:calc(var(--page-padding) * .58) var(--page-padding) calc(var(--page-padding) * .45);border-bottom:2px solid var(--accent);position:relative;text-align:center}.resume-preview header:after{content:'';position:absolute;left:50%;transform:translateX(-50%);bottom:-2px;width:54px;height:4px;border-radius:4px;background:var(--accent)}
.resume-preview header.header-band{background:linear-gradient(135deg,var(--accent),color-mix(in srgb,var(--accent) 72%,#111827));color:#fff;border:0;padding:calc(var(--page-padding) * .58) var(--page-padding)}.resume-preview header.header-band:after{display:none}.resume-preview header.header-inline{display:block}
.resume-preview h1{font-size:2.22em;line-height:1.08;letter-spacing:.08em;color:var(--accent);margin:0;font-weight:800}.resume-preview .header-band h1{color:#fff}
.resume-preview .resume-photo{width:76px;height:96px;margin:0 auto 9px;border:2px solid color-mix(in srgb,var(--accent) 55%,white);border-radius:8px;overflow:hidden;background:color-mix(in srgb,var(--accent) 7%,white);display:grid;place-items:center;color:var(--muted);font-size:.72em;box-shadow:0 4px 12px rgba(15,23,42,.1)}.resume-preview .resume-photo img{width:100%;height:100%;object-fit:cover}.resume-preview .header-band .resume-photo{border-color:rgba(255,255,255,.82);background:rgba(255,255,255,.14);color:rgba(255,255,255,.82)}.resume-preview .role{font-size:1.03em;font-weight:650;letter-spacing:.03em;margin:7px 0 8px}.resume-preview .contact{font-size:.82em;color:var(--muted);overflow-wrap:anywhere;letter-spacing:.015em;display:flex;justify-content:center;gap:4px 8px;flex-wrap:wrap}.resume-preview .header-band .contact{color:rgba(255,255,255,.84)}
.resume-preview .resume-body{padding:8px var(--page-padding) var(--page-padding)}.resume-preview section{break-inside:avoid;margin-top:var(--section-spacing)}.resume-preview section h3{display:flex;align-items:center;gap:9px;font-size:1.08em;line-height:1.35;color:var(--accent);margin:0 0 8px;padding:0 0 5px;border-bottom:1px solid color-mix(in srgb,var(--accent) 42%,white);font-weight:750;letter-spacing:.08em}.resume-preview section h3:before{content:'';width:4px;height:1.05em;border-radius:4px;background:var(--accent);flex:0 0 auto}
.resume-preview .section-list p{position:relative;margin:4px 0;padding-left:13px;white-space:pre-wrap;overflow-wrap:anywhere}.resume-preview .section-list p:before{content:'•';position:absolute;left:0;color:var(--accent)}.resume-preview .section-copy p{margin:0;white-space:pre-wrap;color:#374151}.resume-preview .section-copy p+p{margin-top:5px}
.resume-preview .resume-entry{position:relative;margin:0 0 11px;padding:0 0 0 14px;border-left:2px solid color-mix(in srgb,var(--accent) 25%,white)}.resume-preview .resume-entry:last-child{margin-bottom:0}.resume-preview .resume-entry:before{content:'';position:absolute;left:-5px;top:5px;width:8px;height:8px;border:2px solid #fff;border-radius:50%;background:var(--accent)}.resume-preview .entry-heading{display:flex;justify-content:space-between;gap:14px;align-items:baseline}.resume-preview .entry-heading strong{font-size:1em;color:#111827}.resume-preview .entry-heading span{font-size:.88em;color:var(--muted);margin-left:8px}.resume-preview .entry-heading time{font-size:.78em;color:var(--muted);white-space:nowrap}.resume-preview .entry-description{display:grid;grid-template-columns:1.25em minmax(0,1fr);gap:2px;margin:2px 0;color:#374151;white-space:pre-wrap}.resume-preview .entry-description:first-of-type{margin-top:4px}.resume-preview .entry-description>span:first-child{color:var(--accent);font-weight:700}.resume-preview .entry-tags{display:flex;gap:4px;flex-wrap:wrap;margin-top:5px}.resume-preview .entry-tags span,.resume-preview .skill-list span{border-radius:999px;background:color-mix(in srgb,var(--accent) 9%,white);color:color-mix(in srgb,var(--accent) 76%,#111827);padding:3px 8px;font-size:.78em}.resume-preview .skill-list{display:flex;gap:6px;flex-wrap:wrap}
.resume-preview.section-plain section h3{border:0;text-transform:uppercase;letter-spacing:.13em}.resume-preview.section-band section h3{border:0;background:color-mix(in srgb,var(--accent) 9%,white);padding:6px 9px;border-radius:4px}.resume-preview.section-band section h3:before{height:.85em}
.resume-preview.section-card section{border:1px solid color-mix(in srgb,var(--accent) 18%,white);border-radius:10px;padding:11px 13px;box-shadow:0 3px 12px rgba(15,23,42,.035)}.resume-preview.section-card section h3{border:0}.resume-preview.section-dark section h3{background:linear-gradient(90deg,var(--accent),color-mix(in srgb,var(--accent) 78%,#111827));color:#fff;border:0;padding:6px 10px;border-radius:2px}.resume-preview.section-dark section h3:before{background:#fff}
.resume-preview.has-timeline .resume-entry{border-left-color:color-mix(in srgb,var(--accent) 45%,white)}.resume-preview.layout-two-column .resume-columns{display:grid;grid-template-columns:34% 66%;min-height:inherit}.resume-preview.layout-two-column .resume-column{padding:8px calc(var(--page-padding) * .55) var(--page-padding)}.resume-preview.layout-two-column .resume-column:first-child{background:linear-gradient(180deg,color-mix(in srgb,var(--accent) 8%,white),#fff);padding-left:calc(var(--page-padding) * .72)}.resume-preview.layout-two-column .resume-column:last-child{padding-right:var(--page-padding)}
.resume-preview.layout-sidebar .resume-shell{display:grid;grid-template-columns:31% 69%;min-height:inherit}.resume-preview.layout-sidebar .resume-aside{background:linear-gradient(180deg,var(--sidebar),color-mix(in srgb,var(--sidebar) 82%,white));padding:var(--page-padding) calc(var(--page-padding) * .5);color:#1e3a5f;position:relative}.resume-preview.layout-sidebar .resume-aside:before{content:'';position:absolute;inset:0 auto 0 0;width:7px;background:var(--accent)}.resume-preview.layout-sidebar .resume-aside h1{font-size:2em;letter-spacing:.05em;text-align:center}.resume-preview.layout-sidebar .resume-aside .role{text-align:center}.resume-preview.layout-sidebar .resume-aside .contact{display:grid;gap:5px;margin-top:12px;color:#334155;text-align:center}.resume-preview.layout-sidebar .resume-aside .resume-photo{width:82px;height:104px}.resume-preview.layout-sidebar .resume-main{padding:calc(var(--page-padding) * .45) var(--page-padding) var(--page-padding)}.resume-preview.layout-sidebar section h3{font-size:1em}.resume-preview.layout-sidebar .resume-aside .skill-list{display:grid}.resume-preview.layout-sidebar .resume-aside .skill-list span{background:rgba(255,255,255,.55)}
.resume-preview.is-compact{--resume-font-size:11.5px;--resume-line-height:1.4;--section-spacing:9px}.resume-preview.is-compact h1{font-size:2em}.resume-preview.is-compact .resume-entry{margin-bottom:7px}.resume-preview.is-compact section h3{margin-bottom:5px}
""".strip()


def _sample_resume() -> dict:
    return normalize_resume_structure({
        "basics": {"name": "林微", "title": "AI 产品设计师", "phone": "138****0000", "email": "hello@example.com", "location": "深圳"},
        "sections": [
            {"key": "summary", "title": "个人简介", "content": "关注 AI 产品体验与复杂业务设计，擅长把模糊需求转化为清晰、可落地的产品方案。"},
            {"key": "experience", "title": "工作经历", "entries": [{"company": "微光科技", "role": "高级产品设计师", "start_date": "2022.06", "end_date": "至今", "description": "负责 AI 工作台与智能招聘产品体验设计。\n推动设计系统落地，核心流程效率提升 32%。"}]},
            {"key": "education", "title": "教育经历", "entries": [{"school": "南方科技大学", "major": "工业设计", "degree": "本科", "start_date": "2016.09", "end_date": "2020.06"}]},
            {"key": "projects", "title": "项目经历", "entries": [{"name": "智能招聘工作台", "role": "产品负责人", "start_date": "2024.01", "end_date": "2025.03", "description": "完成从岗位洞察、简历定制到沟通辅助的闭环设计。"}]},
            {"key": "skills", "title": "专业技能", "entries": [{"content": "Figma / 用户研究 / 产品策略 / AI Agent / 数据分析"}]},
            {"key": "evaluation", "title": "自我评价", "content": "有同理心，也有把复杂问题拆解清楚并推进落地的耐心。"},
        ],
    })


def _render_entry(section_key: str, entry: dict) -> str:
    dates = html.escape(_date_range(entry))
    if section_key == "education":
        title = str(entry.get("school") or "学校")
        title_field = "school"
        subtitle_parts = [(field, str(entry.get(field) or "").strip()) for field in ("major", "degree") if entry.get(field)]
    elif section_key == "experience":
        title = str(entry.get("company") or "公司")
        title_field = "company"
        subtitle_parts = [("role", str(entry.get("role") or "").strip())] if entry.get("role") else []
    else:
        title = str(entry.get("name") or "项目名称")
        title_field = "name"
        subtitle_parts = [("role", str(entry.get("role") or "").strip())] if entry.get("role") else []
    subtitle_html = "".join(
        f'<span{_field_css(entry, field)}>{"· " if index else ""}{html.escape(value)}</span>'
        for index, (field, value) in enumerate(subtitle_parts)
    )
    descriptions = "".join(
        f'<p class="entry-description"{_field_css(entry, "description")}><span>{html.escape(marker)}</span><span>{html.escape(text)}</span></p>'
        for _, marker, text in _description_items(entry.get("description") or "")
    )
    technologies = entry.get("technologies") or []
    if isinstance(technologies, str):
        technologies = re.split(r"[,，/、]", technologies)
    tech_html = ""
    if technologies:
        tech_html = '<div class="entry-tags">' + "".join(f'<span{_field_css(entry, "technologies")}>{html.escape(str(value).strip())}</span>' for value in technologies if str(value).strip()) + "</div>"
    return (
        '<article class="resume-entry">'
        f'<div class="entry-heading"><div><strong{_field_css(entry, title_field)}>{html.escape(title)}</strong>{subtitle_html}</div>{f"<time>{dates}</time>" if dates else ""}</div>'
        f"{descriptions}{tech_html}</article>"
    )


def _render_resume_sections(sections: list[dict]) -> str:
    rendered = []
    for section in sections:
        key = str(section.get("key") or "other")
        content = f'<section data-section="{html.escape(key)}"><h3><span>{html.escape(str(section.get("title") or "其他"))}</span></h3>'
        entries = [entry for entry in (section.get("entries") or []) if isinstance(entry, dict)]
        if key in {"education", "experience", "projects"} and entries:
            content += "".join(_render_entry(key, entry) for entry in entries)
        elif key == "skills":
            tokens = []
            for entry in entries:
                tokens.extend((token.strip(), entry) for token in re.split(r"[,，/、|]", str(entry.get("content") or "")) if token.strip())
            content += '<div class="skill-list">' + "".join(f'<span{_field_css(entry, "content")}>{html.escape(token)}</span>' for token, entry in tokens) + "</div>"
        elif key in {"summary", "evaluation"}:
            content += f'<div class="section-copy"{_field_css(section, "content")}>' + "".join(f"<p>{html.escape(item)}</p>" for item in _section_display_items(section)) + "</div>"
        else:
            content += '<div class="section-list">' + "".join(f"<p>{html.escape(item)}</p>" for item in _section_display_items(section)) + "</div>"
        rendered.append(content + "</section>")
    return "".join(rendered)


def template_preview_html(template_id: str, structured: dict | None = None) -> str:
    template = get_template(template_id)
    data = normalize_resume_structure(structured or _sample_resume())
    basics = data.get("basics") or {}
    resume_style = data.get("style") or {}
    accent = resume_style.get("accent_color") or template["accent"]
    muted = template["muted"]
    layout = template.get("layout", "single")
    header_style = template.get("header_style", "plain")
    section_style = template.get("section", "rule")
    ordered = _ordered_sections(data, template)
    basic_visible = "basic" not in set(data.get("hidden_sections") or [])
    name = html.escape(str(basics.get("name") or "个人简历"))
    role = html.escape(str(basics.get("title") or "目标岗位"))
    contact_values = _contact_values(basics)
    contact = " · ".join(html.escape(value) for value in contact_values)
    classes = ["resume-preview", f"layout-{layout}", f"section-{section_style}"]
    if template.get("timeline"):
        classes.append("has-timeline")
    if template.get("compact"):
        classes.append("is-compact")
    style = (
        f"--accent:{accent};--muted:{muted};--sidebar:{template.get('sidebar_bg', '#DBEAFE')};"
        f"--resume-font:'{html.escape(str(resume_style.get('font_family') or template['font']))}';"
        f"--resume-font-size:{resume_style.get('font_size', 13)}px;--resume-line-height:{resume_style.get('line_height', 1.55)};"
        f"--page-padding:{resume_style.get('page_padding', 42)}px;--section-spacing:{resume_style.get('section_spacing', 15)}px"
    )
    header_classes = [f"header-{header_style}", "align-center"]
    header = f'<header class="{" ".join(header_classes)}">{_photo_html(basics)}<h1{_field_css(basics, "name")}>{name}</h1><div class="role"{_field_css(basics, "title")}>{role}</div><div class="contact">{contact}</div></header>' if basic_visible else ""

    if layout == "sidebar":
        side_keys = {"skills", "education", "certifications"}
        side_sections = [section for section in ordered if section.get("key") in side_keys]
        main_sections = [section for section in ordered if section.get("key") not in side_keys]
        aside_contact = "".join(f"<span>{html.escape(value)}</span>" for value in contact_values)
        profile = f'{_photo_html(basics)}<h1{_field_css(basics, "name")}>{name}</h1><div class="role"{_field_css(basics, "title")}>{role}</div><div class="contact">{aside_contact}</div>' if basic_visible else ""
        body = (
            '<div class="resume-shell"><aside class="resume-aside">'
            f'{profile}{_render_resume_sections(side_sections)}'
            f'</aside><main class="resume-main">{_render_resume_sections(main_sections)}</main></div>'
        )
    elif layout == "two-column":
        side_keys = {"summary", "skills", "education", "certifications"}
        left_sections = [section for section in ordered if section.get("key") in side_keys]
        right_sections = [section for section in ordered if section.get("key") not in side_keys]
        body = header + f'<div class="resume-columns"><div class="resume-column">{_render_resume_sections(left_sections)}</div><div class="resume-column">{_render_resume_sections(right_sections)}</div></div>'
    else:
        body = header + f'<main class="resume-body">{_render_resume_sections(ordered)}</main>'
    return f'<div class="{" ".join(classes)}" data-template="{html.escape(template_id)}" style="{style}">{body}</div>'


def _set_docx_font(run, font_name: str, size=None, color=None, bold=None):
    from docx.oxml.ns import qn

    run.font.name = font_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = size
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def _shade_docx_paragraph(paragraph, fill: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.lstrip("#"))
    paragraph._p.get_or_add_pPr().append(shd)


def build_resume_docx(structured: dict, template_id: str = "ats_classic") -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Mm, Pt, RGBColor

    data = normalize_resume_structure(structured)
    template = get_template(template_id)
    resume_style = data.get("style") or {}
    accent_hex = resume_style.get("accent_color") or template["accent"]
    font_name = resume_style.get("font_family") or template["font"]
    base_font_pt = float(resume_style.get("font_size") or 13) * 0.76
    line_height = float(resume_style.get("line_height") or 1.55)
    page_margin = max(0.48, min(0.82, float(resume_style.get("page_padding") or 42) / 96 + 0.2))
    accent = RGBColor.from_string(accent_hex.lstrip("#"))
    muted = RGBColor.from_string(template["muted"].lstrip("#"))
    white = RGBColor(255, 255, 255)
    compact = bool(template.get("compact"))
    band_header = template.get("header_style") == "band"
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Mm(210), Mm(297)
    section.top_margin = section.bottom_margin = Inches(0.55 if compact else page_margin)
    section.left_margin = section.right_margin = Inches(page_margin)
    section.header_distance = section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(9.2 if compact else base_font_pt)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    normal.paragraph_format.space_after = Pt(2.5 if compact else 4)
    normal.paragraph_format.line_spacing = 1.08 if compact else max(1.05, min(1.35, line_height - 0.3))
    for style_name, size in (("Heading 1", 12.5), ("Heading 2", 11.5), ("Heading 3", 10.5)):
        style = doc.styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = accent
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style.paragraph_format.space_before = Pt(7 if compact else 10)
        style.paragraph_format.space_after = Pt(3 if compact else 5)
        style.paragraph_format.keep_with_next = True

    basics = data.get("basics") or {}
    basic_visible = "basic" not in set(data.get("hidden_sections") or [])
    photo_payload = _framed_photo_bytes(basics)
    default_body_line = 1.08 if compact else max(1.05, min(1.35, line_height - 0.3))

    def docx_field(owner, key, default_size_pt, default_line_height):
        size_px, field_line_height = _field_style_values(owner, key, default_size_pt / 0.76, default_line_height)
        return Pt(size_px * 0.76), field_line_height

    def add_docx_section_content(container, section_data, in_sidebar=False):
        key = section_data.get("key")
        entries = [entry for entry in (section_data.get("entries") or []) if isinstance(entry, dict)]
        body_size_pt = 8.7 if in_sidebar else (9.2 if compact else base_font_pt)
        if key in {"education", "experience", "projects"} and entries:
            for entry in entries:
                if key == "education":
                    title_text = str(entry.get("school") or "学校")
                    title_field = "school"
                    subtitle_parts = [(field, str(entry.get(field) or "").strip()) for field in ("major", "degree") if entry.get(field)]
                elif key == "experience":
                    title_text = str(entry.get("company") or "公司")
                    title_field = "company"
                    subtitle_parts = [("role", str(entry.get("role") or "").strip())] if entry.get("role") else []
                else:
                    title_text = str(entry.get("name") or "项目名称")
                    title_field = "name"
                    subtitle_parts = [("role", str(entry.get("role") or "").strip())] if entry.get("role") else []
                title_size, title_line = docx_field(entry, title_field, body_size_pt, default_body_line)
                subtitle_values = [(field, value, *docx_field(entry, field, body_size_pt, default_body_line)) for field, value in subtitle_parts]
                meta = container.add_paragraph()
                meta.paragraph_format.space_before = Pt(3)
                meta.paragraph_format.space_after = Pt(2)
                meta.paragraph_format.keep_with_next = True
                meta.paragraph_format.line_spacing = max([title_line, *[item[3] for item in subtitle_values]])
                _set_docx_font(meta.add_run(title_text), font_name, title_size, RGBColor(17, 24, 39), True)
                for index, (_, value, subtitle_size, _) in enumerate(subtitle_values):
                    _set_docx_font(meta.add_run(f"  {'· ' if index else ''}{value}"), font_name, subtitle_size, muted)
                dates = _date_range(entry)
                if dates:
                    _set_docx_font(meta.add_run(f"    {dates}"), font_name, Pt(max(7.5, body_size_pt - 1)), muted)
                description_size, description_line = docx_field(entry, "description", body_size_pt, default_body_line)
                for line in str(entry.get("description") or "").splitlines():
                    if not line.strip():
                        continue
                    description = _description_items(line)
                    if not description:
                        continue
                    kind, _, text = description[0]
                    paragraph = container.add_paragraph(style="List Number" if kind == "number" else "List Bullet")
                    paragraph.paragraph_format.left_indent = Inches(0.18)
                    paragraph.paragraph_format.first_line_indent = Inches(-0.12)
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.line_spacing = description_line
                    _set_docx_font(paragraph.add_run(text), font_name, description_size, RGBColor(55, 65, 81))
                technologies = entry.get("technologies") or []
                if isinstance(technologies, str):
                    technologies = re.split(r"[,，/、]", technologies)
                if technologies:
                    technology_size, technology_line = docx_field(entry, "technologies", max(7.5, body_size_pt - 1), default_body_line)
                    tech = container.add_paragraph()
                    tech.paragraph_format.space_after = Pt(2)
                    tech.paragraph_format.line_spacing = technology_line
                    _set_docx_font(tech.add_run("技术：" + " / ".join(str(value).strip() for value in technologies if str(value).strip())), font_name, technology_size, accent)
            return
        if key == "skills" and entries:
            for entry in entries:
                value = str(entry.get("content") or "").strip()
                if not value:
                    continue
                field_size, field_line = docx_field(entry, "content", body_size_pt, default_body_line)
                paragraph = container.add_paragraph(style="List Bullet")
                paragraph.paragraph_format.left_indent = Inches(0.18)
                paragraph.paragraph_format.first_line_indent = Inches(-0.12)
                paragraph.paragraph_format.space_after = Pt(2.5)
                paragraph.paragraph_format.line_spacing = field_line
                _set_docx_font(paragraph.add_run(value), font_name, field_size, RGBColor(55, 65, 81))
            return
        items = _section_display_items(section_data)
        plain_text = key in {"summary", "evaluation"}
        field_size, field_line = docx_field(section_data, "content", body_size_pt, default_body_line)
        for item in items:
            paragraph = container.add_paragraph() if plain_text else container.add_paragraph(style="List Bullet")
            if not plain_text:
                paragraph.paragraph_format.left_indent = Inches(0.18)
                paragraph.paragraph_format.first_line_indent = Inches(-0.12)
            paragraph.paragraph_format.space_after = Pt(2.5)
            paragraph.paragraph_format.line_spacing = field_line
            _set_docx_font(paragraph.add_run(str(item).strip()), font_name, field_size, RGBColor(55, 65, 81))

    def add_docx_cell_section(cell, section_data, in_side_column=False):
        first_paragraph = cell.paragraphs[0]
        heading = first_paragraph if len(cell.paragraphs) == 1 and not first_paragraph.text else cell.add_paragraph()
        heading.paragraph_format.space_before = Pt(7 if in_side_column else 9)
        heading.paragraph_format.space_after = Pt(4)
        heading.paragraph_format.keep_with_next = True
        heading_color = RGBColor(30, 58, 95) if in_side_column else accent
        _set_docx_font(heading.add_run(str(section_data.get("title") or "其他")), font_name, Pt(11 if in_side_column else 11.5), heading_color, True)
        if not in_side_column:
            p_pr = heading._p.get_or_add_pPr()
            borders = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            for key, value in (("val", "single"), ("sz", "5"), ("space", "2"), ("color", accent_hex.lstrip("#"))):
                bottom.set(qn(f"w:{key}"), value)
            borders.append(bottom)
            p_pr.append(borders)
        add_docx_section_content(cell, section_data, in_side_column)

    if template.get("layout") == "sidebar":
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        left_cell, main_cell = table.rows[0].cells
        usable_width = section.page_width.inches - page_margin * 2
        left_width, main_width = usable_width * 0.3, usable_width * 0.7
        left_cell.width, main_cell.width = Inches(left_width), Inches(main_width)
        table.columns[0].width, table.columns[1].width = Inches(left_width), Inches(main_width)
        _shade_docx_paragraph(left_cell.paragraphs[0], template.get("sidebar_bg", "DBEAFE"))
        cell_shading = OxmlElement("w:shd")
        cell_shading.set(qn("w:fill"), template.get("sidebar_bg", "#DBEAFE").lstrip("#"))
        left_cell._tc.get_or_add_tcPr().append(cell_shading)
        for cell in (left_cell, main_cell):
            margins = OxmlElement("w:tcMar")
            for edge, value in (("top", "150"), ("start", "150"), ("bottom", "150"), ("end", "150")):
                node = OxmlElement(f"w:{edge}")
                node.set(qn("w:w"), value)
                node.set(qn("w:type"), "dxa")
                margins.append(node)
            cell._tc.get_or_add_tcPr().append(margins)

        if basic_visible:
            sidebar_name_size, sidebar_name_line = docx_field(basics, "name", 21, 1.08)
            sidebar_role_size, sidebar_role_line = docx_field(basics, "title", 10, 1.35)
            name_p = left_cell.paragraphs[0]
            name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_p.paragraph_format.space_after = Pt(2)
            name_p.paragraph_format.line_spacing = sidebar_name_line
            name_p.add_run().add_picture(io.BytesIO(photo_payload), width=Inches(0.76), height=Inches(1.01))
            name_p.add_run().add_break()
            _set_docx_font(name_p.add_run(str(basics.get("name") or "个人简历")), font_name, sidebar_name_size, accent, True)
            role_p = left_cell.add_paragraph()
            role_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            role_p.paragraph_format.space_after = Pt(8)
            role_p.paragraph_format.line_spacing = sidebar_role_line
            _set_docx_font(role_p.add_run(str(basics.get("title") or "目标岗位")), font_name, sidebar_role_size, RGBColor(30, 58, 95), True)
            contact_values = _contact_values(basics)
            for value in contact_values:
                contact_p = left_cell.add_paragraph()
                contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                contact_p.paragraph_format.space_after = Pt(2)
                _set_docx_font(contact_p.add_run(value), font_name, Pt(8), RGBColor(51, 65, 85))

        side_keys = {"skills", "education", "certifications"}
        ordered = _ordered_sections(data, template)

        for section_data in ordered:
            destination = left_cell if section_data.get("key") in side_keys else main_cell
            add_docx_cell_section(destination, section_data, destination is left_cell)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_docx_font(footer.add_run("由 MyJob JD 定制简历生成"), font_name, Pt(7.5), muted)
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()

    if basic_visible:
        header_name_size, header_name_line = docx_field(basics, "name", 22 if compact else 25, 1.08)
        header_title_size, header_title_line = docx_field(basics, "title", 10.5, 1.35)
        photo_paragraph = doc.add_paragraph()
        photo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        photo_paragraph.paragraph_format.space_after = Pt(3)
        if band_header:
            photo_paragraph.paragraph_format.space_before = Pt(7)
            _shade_docx_paragraph(photo_paragraph, accent_hex)
        photo_paragraph.add_run().add_picture(io.BytesIO(photo_payload), width=Inches(0.76), height=Inches(1.01))
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.paragraph_format.space_after = Pt(0 if band_header else 1)
        header.paragraph_format.line_spacing = header_name_line
        if band_header:
            header.paragraph_format.left_indent = header.paragraph_format.right_indent = Inches(0.12)
            header.paragraph_format.space_before = Pt(7)
            _shade_docx_paragraph(header, accent_hex)
        name_run = header.add_run(str(basics.get("name") or "个人简历"))
        _set_docx_font(name_run, font_name, header_name_size, white if band_header else accent, True)
        if basics.get("title"):
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.paragraph_format.space_after = Pt(0 if band_header else 3)
            title.paragraph_format.line_spacing = header_title_line
            if band_header:
                title.paragraph_format.left_indent = title.paragraph_format.right_indent = Inches(0.12)
                _shade_docx_paragraph(title, accent_hex)
            _set_docx_font(title.add_run(str(basics["title"])), font_name, header_title_size, white if band_header else muted, True)
        contact_values = _contact_values(basics)
        if contact_values:
            contact = doc.add_paragraph()
            contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact.paragraph_format.space_after = Pt(7)
            if band_header:
                contact.paragraph_format.left_indent = contact.paragraph_format.right_indent = Inches(0.12)
                contact.paragraph_format.space_after = Pt(8)
                _shade_docx_paragraph(contact, accent_hex)
            _set_docx_font(contact.add_run("  |  ".join(contact_values)), font_name, Pt(8.5), white if band_header else muted)

    if template.get("layout") == "two-column":
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        side_cell, main_cell = table.rows[0].cells
        usable_width = section.page_width.inches - page_margin * 2
        side_width, main_width = usable_width * 0.34, usable_width * 0.66
        side_cell.width, main_cell.width = Inches(side_width), Inches(main_width)
        table.columns[0].width, table.columns[1].width = Inches(side_width), Inches(main_width)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F4F8FC")
        side_cell._tc.get_or_add_tcPr().append(shading)
        for cell in (side_cell, main_cell):
            margins = OxmlElement("w:tcMar")
            for edge, value in (("top", "110"), ("start", "140"), ("bottom", "110"), ("end", "140")):
                node = OxmlElement(f"w:{edge}")
                node.set(qn("w:w"), value)
                node.set(qn("w:type"), "dxa")
                margins.append(node)
            cell._tc.get_or_add_tcPr().append(margins)
        side_keys = {"summary", "skills", "education"}
        for section_data in _ordered_sections(data, template):
            destination = side_cell if section_data.get("key") in side_keys else main_cell
            add_docx_cell_section(destination, section_data, destination is side_cell)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_docx_font(footer.add_run("由 MyJob JD 定制简历生成"), font_name, Pt(7.5), muted)
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()

    for section_data in _ordered_sections(data, template):
        heading = doc.add_paragraph(style="Heading 1")
        heading.add_run(str(section_data.get("title") or "其他"))
        if template["section"] == "rule":
            p_pr = heading._p.get_or_add_pPr()
            borders = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            for key, value in (("val", "single"), ("sz", "5"), ("space", "2"), ("color", accent_hex.lstrip("#"))):
                bottom.set(qn(f"w:{key}"), value)
            borders.append(bottom)
            p_pr.append(borders)
        elif template["section"] in {"band", "card"}:
            _shade_docx_paragraph(heading, "EEF2F7")
        elif template["section"] == "dark":
            _shade_docx_paragraph(heading, accent_hex)
            for run in heading.runs:
                run.font.color.rgb = white
        add_docx_section_content(doc, section_data)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_docx_font(footer.add_run("由 MyJob JD 定制简历生成"), font_name, Pt(7.5), muted)
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def build_resume_pdf(structured: dict, template_id: str = "ats_classic") -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import HRFlowable, Image as ReportLabImage, KeepInFrame, ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    data = normalize_resume_structure(structured)
    template = get_template(template_id)
    resume_style = data.get("style") or {}
    accent_hex = resume_style.get("accent_color") or template["accent"]
    font_size = float(resume_style.get("font_size") or 13)
    line_height = float(resume_style.get("line_height") or 1.55)
    page_margin = max(0.48, min(0.82, float(resume_style.get("page_padding") or 42) / 96 + 0.2))
    regular_name, bold_name = "STSong-Light", "STSong-Light"
    font_candidates = [
        (Path("C:/Windows/Fonts/Deng.ttf"), Path("C:/Windows/Fonts/Dengb.ttf")),
        (Path("C:/Windows/Fonts/NotoSansSC-VF.ttf"), Path("C:/Windows/Fonts/NotoSansSC-VF.ttf")),
        (Path("C:/Windows/Fonts/simhei.ttf"), Path("C:/Windows/Fonts/simhei.ttf")),
        (Path("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttf"), Path("/usr/share/fonts/truetype/noto/NotoSansCJK-Bold.ttf")),
    ]
    for regular_path, bold_path in font_candidates:
        if regular_path.exists():
            regular_name, bold_name = "MyJobResumeCN", "MyJobResumeCNBold"
            if regular_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(regular_name, str(regular_path)))
            if bold_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(bold_name, str(bold_path if bold_path.exists() else regular_path)))
            break
    else:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    buffer = io.BytesIO()
    compact = bool(template.get("compact"))
    layout = template.get("layout", "single")
    sidebar_width_inches = 2.35
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=page_margin * inch,
        leftMargin=(sidebar_width_inches + 0.32 if layout == "sidebar" else page_margin) * inch,
        topMargin=(0.55 if compact else page_margin) * inch,
        bottomMargin=(0.55 if compact else page_margin) * inch,
        title=str((data.get("basics") or {}).get("name") or "个人简历"),
        author="MyJob",
    )
    accent = colors.HexColor(accent_hex)
    muted = colors.HexColor(template["muted"])
    band_header = template.get("header_style") == "band"
    styles = getSampleStyleSheet()
    align = TA_CENTER
    basics = data.get("basics") or {}
    default_name_pt = 22 if compact else 25
    name_px, name_line = _field_style_values(basics, "name", default_name_pt / 0.76, 27 / default_name_pt)
    role_px, role_line = _field_style_values(basics, "title", 10.5 / 0.76, 14 / 10.5)
    name_pt, role_pt = name_px * 0.76, role_px * 0.76
    name_style = ParagraphStyle("ResumeName", parent=styles["Title"], fontName=bold_name, fontSize=name_pt, leading=name_pt * name_line, textColor=colors.white if band_header else accent, alignment=align, spaceAfter=2)
    role_style = ParagraphStyle("ResumeRole", parent=styles["Normal"], fontName=bold_name, fontSize=role_pt, leading=role_pt * role_line, textColor=colors.white if band_header else muted, alignment=align, spaceAfter=2)
    contact_style = ParagraphStyle("ResumeContact", parent=styles["Normal"], fontName=regular_name, fontSize=8.5, leading=12, textColor=colors.HexColor("#E5E7EB") if band_header else muted, alignment=align, spaceAfter=7)
    heading_style = ParagraphStyle("ResumeHeading", parent=styles["Heading2"], fontName=bold_name, fontSize=12.5, leading=16, textColor=accent, spaceBefore=7 if compact else 10, spaceAfter=3)
    body_pt = 9.2 if compact else font_size * 0.76
    if layout in {"two-column", "sidebar"}:
        body_pt = max(8.6, body_pt - 0.3)
    body_style = ParagraphStyle("ResumeBody", parent=styles["Normal"], fontName=regular_name, fontSize=body_pt, leading=body_pt * max(1.15, min(1.55, line_height)), textColor=colors.HexColor("#202124"), spaceAfter=2, wordWrap="CJK")
    entry_style = ParagraphStyle("ResumeEntry", parent=body_style, fontName=regular_name, leading=body_style.leading, spaceAfter=1)
    date_style = ParagraphStyle("ResumeDate", parent=body_style, fontName=regular_name, fontSize=max(7.5, body_pt - 1), textColor=muted, alignment=TA_RIGHT, spaceAfter=1)
    field_style_counter = [0]

    def pdf_field_style(owner, key, parent, default_size_pt=None, default_line=None, **overrides):
        default_size_pt = float(default_size_pt or parent.fontSize)
        default_line = float(default_line or (parent.leading / parent.fontSize))
        size_px, selected_line = _field_style_values(owner, key, default_size_pt / 0.76, default_line)
        size_pt = size_px * 0.76
        field_style_counter[0] += 1
        return ParagraphStyle(
            f"ResumeField{field_style_counter[0]}",
            parent=parent,
            fontSize=size_pt,
            leading=size_pt * selected_line,
            **overrides,
        )

    basic_visible = "basic" not in set(data.get("hidden_sections") or [])
    photo_payload = _framed_photo_bytes(basics)
    photo_flowable = ReportLabImage(io.BytesIO(photo_payload), width=0.76 * inch, height=1.01 * inch)
    photo_holder = Table([[photo_flowable]], colWidths=[0.82 * inch])
    photo_holder.hAlign = "CENTER"
    photo_holder.setStyle(TableStyle([
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
    ]))
    header_story = [photo_holder, Spacer(1, 4), Paragraph(html.escape(str(basics.get("name") or "个人简历")), name_style)] if basic_visible else []
    if basic_visible and basics.get("title"):
        header_story.append(Paragraph(html.escape(str(basics["title"])), role_style))
    contact = _contact_values(basics)
    if basic_visible and contact:
        header_story.append(Paragraph(html.escape("  |  ".join(contact)), contact_style))
    if layout == "sidebar":
        story = []
    elif band_header and header_story:
        header_table = Table([[item] for item in header_story], colWidths=[doc.width])
        header_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), accent),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 14),
            ("RIGHTPADDING", (0, 0), (-1, -1), 14),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (0, 0), 10),
            ("BOTTOMPADDING", (0, -1), (0, -1), 5),
        ]))
        story = [header_table, Spacer(1, 7)]
    else:
        story = header_story
    def section_flowables(section_data, content_width):
        result = []
        heading = Paragraph(html.escape(str(section_data.get("title") or "其他")), heading_style)
        if template["section"] in {"band", "card", "dark"}:
            heading_bg = accent if template["section"] == "dark" else colors.HexColor("#EEF2F7")
            if template["section"] == "dark":
                dark_heading = ParagraphStyle("ResumeHeadingDark", parent=heading_style, textColor=colors.white, spaceBefore=0, spaceAfter=0)
                heading = Paragraph(html.escape(str(section_data.get("title") or "其他")), dark_heading)
            heading_table = Table([[heading]], colWidths=[content_width])
            heading_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), heading_bg),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            result.append(heading_table)
        else:
            result.append(heading)
        if template["section"] == "rule":
            result.append(HRFlowable(width="100%", thickness=0.6, color=accent, spaceBefore=0, spaceAfter=4))
        key = section_data.get("key")
        entries = [entry for entry in (section_data.get("entries") or []) if isinstance(entry, dict)]
        if key in {"education", "experience", "projects"} and entries:
            for entry in entries:
                if key == "education":
                    title_text = str(entry.get("school") or "学校")
                    title_field = "school"
                    subtitle_parts = [(field, str(entry.get(field) or "").strip()) for field in ("major", "degree") if entry.get(field)]
                elif key == "experience":
                    title_text = str(entry.get("company") or "公司")
                    title_field = "company"
                    subtitle_parts = [("role", str(entry.get("role") or "").strip())] if entry.get("role") else []
                else:
                    title_text = str(entry.get("name") or "项目名称")
                    title_field = "name"
                    subtitle_parts = [("role", str(entry.get("role") or "").strip())] if entry.get("role") else []
                title_style = pdf_field_style(entry, title_field, entry_style)
                subtitle_styles = [(field, value, pdf_field_style(entry, field, entry_style)) for field, value in subtitle_parts]
                meta_style = ParagraphStyle(
                    f"ResumeEntryMeta{field_style_counter[0]}",
                    parent=title_style,
                    leading=max([title_style.leading, *[item[2].leading for item in subtitle_styles]]),
                )
                headline = f'<b><font size="{title_style.fontSize:g}">{html.escape(title_text)}</font></b>'
                for index, (_, value, subtitle_style) in enumerate(subtitle_styles):
                    headline += f'  <font color="{template["muted"]}" size="{subtitle_style.fontSize:g}">{"· " if index else ""}{html.escape(value)}</font>'
                meta_table = Table(
                    [[Paragraph(headline, meta_style), Paragraph(html.escape(_date_range(entry)), date_style)]],
                    colWidths=[content_width * 0.72, content_width * 0.28],
                )
                meta_table.setStyle(TableStyle([
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
                ]))
                result.append(meta_table)
                description_items = _description_items(entry.get("description") or "")
                groups = []
                for kind, marker, text in description_items:
                    if not groups or groups[-1][0] != kind:
                        groups.append([kind, marker, []])
                    groups[-1][2].append(text)
                for kind, marker, group_items in groups:
                    start = int(marker.rstrip(".")) if kind == "number" and marker.rstrip(".").isdigit() else "circle"
                    description_style = pdf_field_style(entry, "description", body_style)
                    result.append(ListFlowable([ListItem(Paragraph(html.escape(line), description_style), leftIndent=10) for line in group_items], bulletType="1" if kind == "number" else "bullet", start=start, leftIndent=14, bulletFontName=regular_name, bulletFontSize=6, spaceAfter=3))
                technologies = entry.get("technologies") or []
                if isinstance(technologies, str):
                    technologies = re.split(r"[,，/、]", technologies)
                if technologies:
                    technology_style = pdf_field_style(entry, "technologies", date_style, alignment=TA_LEFT)
                    result.append(Paragraph("技术：" + html.escape(" / ".join(str(value).strip() for value in technologies if str(value).strip())), technology_style))
        else:
            items = _section_display_items(section_data)
            if key in {"summary", "evaluation"}:
                section_body_style = pdf_field_style(section_data, "content", body_style)
                result.extend(Paragraph(html.escape(item), section_body_style) for item in items)
            elif key == "skills" and entries:
                for entry in entries:
                    value = str(entry.get("content") or "").strip()
                    if not value:
                        continue
                    skill_style = pdf_field_style(entry, "content", body_style)
                    result.append(ListFlowable([ListItem(Paragraph(html.escape(value), skill_style), leftIndent=10)], bulletType="bullet", start="circle", leftIndent=14, bulletFontName=regular_name, bulletFontSize=6, spaceAfter=2))
            else:
                flow_items = [ListItem(Paragraph(html.escape(item), body_style), leftIndent=10) for item in items]
                if flow_items:
                    result.append(ListFlowable(flow_items, bulletType="bullet", start="circle", leftIndent=14, bulletFontName=regular_name, bulletFontSize=6, spaceAfter=2))
        return result

    ordered_sections = _ordered_sections(data, template)
    sidebar_section_keys = {"skills", "education"}
    sidebar_sections = [section_data for section_data in ordered_sections if section_data.get("key") in sidebar_section_keys]
    if layout == "two-column":
        side_keys = {"summary", "skills", "education"}
        side_width, main_width = doc.width * 0.34, doc.width * 0.66
        side_content = [item for section_data in ordered_sections if section_data.get("key") in side_keys for item in section_flowables(section_data, side_width - 12)]
        main_content = [item for section_data in ordered_sections if section_data.get("key") not in side_keys for item in section_flowables(section_data, main_width - 12)]
        available_height = max(3 * inch, doc.height - (0.92 * inch if header_story else 0.2 * inch))
        columns = Table(
            [[KeepInFrame(side_width - 12, available_height, side_content, mode="shrink"), KeepInFrame(main_width - 12, available_height, main_content, mode="shrink")]],
            colWidths=[side_width, main_width],
        )
        columns.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#F4F8FC")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        story.append(columns)
    else:
        for section_data in ordered_sections:
            if layout == "sidebar" and section_data.get("key") in sidebar_section_keys:
                continue
            story.extend(section_flowables(section_data, doc.width))
    if layout == "sidebar":
        sidebar_bg = colors.HexColor(template.get("sidebar_bg", "#DBEAFE"))

        def draw_sidebar(canvas, _doc):
            page_width, page_height = A4
            sidebar_width = sidebar_width_inches * inch
            canvas.saveState()
            canvas.setFillColor(sidebar_bg)
            canvas.rect(0, 0, sidebar_width, page_height, stroke=0, fill=1)
            canvas.setFillColor(accent)
            canvas.rect(0, page_height - 0.16 * inch, sidebar_width, 0.16 * inch, stroke=0, fill=1)
            side_name_base = ParagraphStyle("SidebarNameBase", fontName=bold_name, fontSize=20, leading=24, textColor=accent, spaceAfter=4, alignment=TA_CENTER)
            side_role_base = ParagraphStyle("SidebarRoleBase", fontName=bold_name, fontSize=10, leading=14, textColor=colors.HexColor("#1E3A5F"), spaceAfter=10, alignment=TA_CENTER)
            side_name = pdf_field_style(basics, "name", side_name_base)
            side_role = pdf_field_style(basics, "title", side_role_base)
            side_contact = ParagraphStyle("SidebarContact", fontName=regular_name, fontSize=8, leading=12, textColor=colors.HexColor("#334155"), wordWrap="CJK", alignment=TA_CENTER)
            side_heading = ParagraphStyle("SidebarHeading", fontName=bold_name, fontSize=10.5, leading=13, textColor=colors.HexColor("#1E3A5F"), spaceBefore=8, spaceAfter=3, wordWrap="CJK")
            side_body = ParagraphStyle("SidebarBody", fontName=regular_name, fontSize=7.7, leading=11, textColor=colors.HexColor("#334155"), spaceAfter=3, wordWrap="CJK")
            photo_width, photo_height = 0.82 * inch, 1.09 * inch
            photo_x = (sidebar_width - photo_width) / 2
            photo_y = page_height - 0.5 * inch - photo_height
            canvas.drawImage(ImageReader(io.BytesIO(photo_payload)), photo_x, photo_y, width=photo_width, height=photo_height, mask="auto", preserveAspectRatio=True)
            y = photo_y - 0.12 * inch
            flowables = [] if not basic_visible else [
                Paragraph(html.escape(str(basics.get("name") or "个人简历")), side_name),
                Paragraph(html.escape(str(basics.get("title") or "目标岗位")), side_role),
                Paragraph("<br/>".join(html.escape(value) for value in contact), side_contact),
            ]
            for section_data in sidebar_sections:
                flowables.append(Paragraph(html.escape(str(section_data.get("title") or "其他")), side_heading))
                entries = [entry for entry in (section_data.get("entries") or []) if isinstance(entry, dict)]
                if section_data.get("key") == "education" and entries:
                    for entry in entries:
                        school = html.escape(str(entry.get("school") or "学校"))
                        school_style = pdf_field_style(entry, "school", side_body)
                        detail_parts = []
                        detail_leading = [school_style.leading]
                        for index, field in enumerate(("major", "degree")):
                            value = str(entry.get(field) or "").strip()
                            if not value:
                                continue
                            value_style = pdf_field_style(entry, field, side_body)
                            detail_leading.append(value_style.leading)
                            detail_parts.append(f'<font size="{value_style.fontSize:g}">{"· " if index and detail_parts else ""}{html.escape(value)}</font>')
                        dates = _date_range(entry)
                        if dates:
                            detail_parts.append(html.escape(dates))
                        school_paragraph_style = ParagraphStyle(f"SidebarEducation{field_style_counter[0]}", parent=school_style, leading=max(detail_leading))
                        detail = "<br/>".join(detail_parts)
                        flowables.append(Paragraph(f'<b><font size="{school_style.fontSize:g}">{school}</font></b>' + (f"<br/>{detail}" if detail else ""), school_paragraph_style))
                elif section_data.get("key") == "skills" and entries:
                    for entry in entries:
                        value = str(entry.get("content") or "").strip()
                        if value:
                            flowables.append(Paragraph(html.escape(value), pdf_field_style(entry, "content", side_body)))
                else:
                    items = _section_display_items(section_data)
                    if items:
                        flowables.append(Paragraph(" / ".join(html.escape(item) for item in items), side_body))
            for flowable in flowables:
                width, height = flowable.wrap(sidebar_width - 0.54 * inch, page_height)
                if y - height < 0.35 * inch:
                    break
                y -= height
                flowable.drawOn(canvas, 0.28 * inch, y)
                y -= 3
            canvas.restoreState()

        doc.build(story, onFirstPage=draw_sidebar, onLaterPages=draw_sidebar)
    else:
        doc.build(story)
    return buffer.getvalue()


def build_resume_html(structured: dict, template_id: str = "ats_classic") -> str:
    preview = template_preview_html(template_id, structured)
    return f'''<!doctype html><html lang="zh-CN"><head><meta charset="utf-8"><title>简历</title><style>
    @page{{size:A4;margin:0}}body{{margin:0;background:#f3f4f6;color:#202124}}{resume_preview_css()}.resume-preview{{width:210mm;min-height:297mm;margin:20px auto}}@media print{{body{{background:white}}.resume-preview{{margin:0;width:210mm;min-height:297mm}}}}</style></head><body>{preview}</body></html>'''


def build_resume_bytes(structured: dict, template_id: str, output_format: str) -> tuple[bytes, str, str]:
    fmt = output_format.lower().strip(".")
    if fmt == "docx":
        return build_resume_docx(structured, template_id), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"
    if fmt == "pdf":
        return build_resume_pdf(structured, template_id), "application/pdf", "pdf"
    if fmt in {"html", "htm"}:
        return build_resume_html(structured, template_id).encode("utf-8"), "text/html; charset=utf-8", "html"
    if fmt in {"md", "markdown", "txt"}:
        suffix = "txt" if fmt == "txt" else "md"
        return structured_to_markdown(structured).encode("utf-8"), "text/plain; charset=utf-8", suffix
    raise ValueError("format 支持 docx、pdf、html、md、txt")
